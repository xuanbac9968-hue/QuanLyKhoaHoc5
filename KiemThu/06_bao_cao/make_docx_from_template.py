#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Tạo BaoCao_KiemThu_NGUYENXUANBAC.docx
Theo đúng hình thức file mẫu H11. Trần Huyền Trang.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re, os, copy

BASE = r'D:\QuanLyKhoaHoc5\KiemThu\06_bao_cao'
OUT  = os.path.join(BASE, 'BaoCao_KiemThu_NGUYENXUANBAC.docx')

# ─────────────────────────────────────────────
# KHỞI TẠO TÀI LIỆU
# ─────────────────────────────────────────────
doc = Document()

# Xóa đoạn mặc định
for p in list(doc.paragraphs):
    p._element.getparent().remove(p._element)

# Normal style – Times New Roman 13pt, ls 1.3, sb/sa 6pt
ns = doc.styles['Normal']
ns.font.name = 'Times New Roman'
ns.font.size = Pt(13)
pf = ns.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
pf.line_spacing = 1.3
pf.space_before = Pt(6)
pf.space_after  = Pt(6)
pf.first_line_indent = None
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# ─────────────────────────────────────────────
# CẤU HÌNH SECTIONS
# ─────────────────────────────────────────────
def _cfg_sec(sec):
    sec.page_width    = Cm(21)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(3.5)
    sec.right_margin  = Cm(2.0)
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.header_distance = Cm(1.27)
    sec.footer_distance = Cm(1.27)

_cfg_sec(doc.sections[0])

# ─────────────────────────────────────────────
# HÀM TIỆN ÍCH FONT
# ─────────────────────────────────────────────
def _sf(run, sz=13, bold=False, italic=False, name='Times New Roman'):
    run.font.name   = name
    run.font.size   = Pt(sz)
    run.font.bold   = bold
    run.font.italic = italic
    rPr = run._r.get_or_add_rPr()
    rF  = rPr.find(qn('w:rFonts'))
    if rF is None:
        rF = OxmlElement('w:rFonts')
        rPr.insert(0, rF)
    for a in ['w:ascii','w:hAnsi','w:eastAsia','w:cs']:
        rF.set(qn(a), name)

def _fmtp(para, align='justify', sb=6, sa=6, fi=None, li=0, ls=1.3):
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing   = ls
    pf.space_before   = Pt(sb)
    pf.space_after    = Pt(sa)
    pf.first_line_indent = Cm(fi) if fi else None
    if li: pf.left_indent = Cm(li)
    para.alignment = {
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
        'center':  WD_ALIGN_PARAGRAPH.CENTER,
        'left':    WD_ALIGN_PARAGRAPH.LEFT,
        'right':   WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)

def _parse(text):
    """Trả về list (text, bold, italic) sau khi parse **bold**, *italic*, `code`"""
    res, last = [], 0
    for m in re.finditer(r'\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`', text):
        if m.start() > last:
            res.append((text[last:m.start()], False, False))
        if   m.group(1): res.append((m.group(1), True, False))
        elif m.group(2): res.append((m.group(2), False, True))
        else:            res.append((m.group(3), False, False))
        last = m.end()
    if last < len(text): res.append((text[last:], False, False))
    return res or [(text, False, False)]

# ─────────────────────────────────────────────
# CÁC HÀM THÊM NỘI DUNG
# ─────────────────────────────────────────────

def P(text='', bold=False, italic=False, align='justify', sz=13,
      sb=6, sa=6, fi=None, li=0, inline=True):
    p = doc.add_paragraph()
    _fmtp(p, align=align, sb=sb, sa=sa, fi=fi, li=li)
    if text:
        if inline:
            for t, b, i in _parse(text):
                r = p.add_run(t); _sf(r, sz, bold or b, italic or i)
        else:
            r = p.add_run(text); _sf(r, sz, bold, italic)
    return p

def H1(text, align='center'):
    """Heading 1 – chương hoặc tiêu đề đặc biệt: 13pt bold center/justify"""
    doc.add_page_break()
    p = doc.add_paragraph(style='Heading 1')
    p.clear()
    _fmtp(p, align=align, sb=6, sa=6, fi=None)
    r = p.add_run(text); _sf(r, 13, True)
    return p

def H1_nobreak(text, align='center'):
    """Heading 1 không có page break trước"""
    p = doc.add_paragraph(style='Heading 1')
    p.clear()
    _fmtp(p, align=align, sb=6, sa=6, fi=None)
    r = p.add_run(text); _sf(r, 13, True)
    return p

def H2(text):
    """Heading 2 – mục 1.x: 13pt NOT bold, justify"""
    p = doc.add_paragraph(style='Heading 2')
    p.clear()
    _fmtp(p, align='justify', sb=6, sa=6, fi=None)
    r = p.add_run(text); _sf(r, 13, False)
    return p

def H3(text):
    """Heading 3 – tiểu mục 1.x.x: 13pt NOT bold, justify"""
    p = doc.add_paragraph(style='Heading 3')
    p.clear()
    _fmtp(p, align='justify', sb=6, sa=6, fi=None)
    r = p.add_run(text); _sf(r, 13, False)
    return p

def LP(text, bullet=True):
    """List Paragraph: indent 1.27cm, no first-line"""
    p = doc.add_paragraph(style='List Paragraph')
    p.clear()
    _fmtp(p, align='justify', sb=0, sa=0, fi=None, li=1.27)
    prefix = '- ' if bullet else ''
    parts = _parse(text)
    first = True
    for t, b, i in parts:
        r = p.add_run((prefix if first else '') + t)
        _sf(r, 13, b, i); first = False
    return p

def LP_bold(text):
    """List Paragraph với bold lead"""
    p = doc.add_paragraph(style='List Paragraph')
    p.clear()
    _fmtp(p, align='justify', sb=0, sa=6, fi=None, li=1.27)
    r = p.add_run(text); _sf(r, 13, True)
    return p

def NUMBERED(n, text):
    p = doc.add_paragraph(style='List Paragraph')
    p.clear()
    _fmtp(p, align='justify', sb=0, sa=0, fi=None, li=1.27)
    parts = _parse(text)
    first = True
    for t, b, i in parts:
        r = p.add_run((f'{n}. ' if first else '') + t)
        _sf(r, 13, b, i); first = False
    return p

def CAPTION_TABLE(text):
    """Chú thích bảng (style bảng): 13pt center"""
    p = doc.add_paragraph()
    _fmtp(p, align='center', sb=6, sa=6, fi=None)
    r = p.add_run(text); _sf(r, 13, False)
    return p

def CAPTION_IMG(text):
    """Chú thích hình (style Ảnh): 13pt center italic"""
    p = doc.add_paragraph()
    _fmtp(p, align='center', sb=6, sa=6, fi=None)
    r = p.add_run(text); _sf(r, 13, False, True)
    return p

def CODE(lines):
    for line in lines:
        p = doc.add_paragraph()
        _fmtp(p, align='left', sb=0, sa=0, fi=None, li=1.0, ls=1.1)
        r = p.add_run(line if line.strip() else ' ')
        r.font.name = 'Courier New'
        r.font.size = Pt(10)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F2F2F2')
        pPr.append(shd)

def QUOTE(text):
    p = doc.add_paragraph()
    _fmtp(p, align='justify', sb=3, sa=3, fi=None, li=0.5)
    r = p.add_run(text); _sf(r, 13, False, True)

def TABLE(headers, rows):
    nc = max(len(headers), max((len(r) for r in rows), default=0))
    if nc == 0: return
    nr = len(rows) + (1 if headers else 0)
    tbl = doc.add_table(rows=nr, cols=nc)
    tbl.style = 'Table Grid'
    r0 = 0
    if headers:
        hr = tbl.rows[0]
        for i, h in enumerate(headers[:nc]):
            c = hr.cells[i]; c.text = ''
            cp = c.paragraphs[0]
            _fmtp(cp, align='center', sb=3, sa=3, fi=None)
            run = cp.add_run(h); _sf(run, 13, True)
            # Nền xanh nhạt
            tcPr = c._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'D9E2F3')
            tcPr.append(shd)
        r0 = 1
    for ri, row_data in enumerate(rows):
        row = tbl.rows[r0+ri]
        for ci, ct in enumerate(row_data[:nc]):
            c = row.cells[ci]; c.text = ''
            cp = c.paragraphs[0]
            _fmtp(cp, align='justify', sb=2, sa=2, fi=None)
            run = cp.add_run(str(ct).strip()); _sf(run, 13)
    # Khoảng cách sau bảng
    gap = doc.add_paragraph(); _fmtp(gap, sb=0, sa=6, fi=None)
    return tbl

def SIGN_TABLE(left_text, right_text):
    """Bảng chữ ký 1x2 như trong mẫu"""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    tbl.columns[0].width = Cm(8)
    tbl.columns[1].width = Cm(7)
    for ci, txt in enumerate([left_text, right_text]):
        c = tbl.rows[0].cells[ci]; c.text = ''
        cp = c.paragraphs[0]
        _fmtp(cp, align='justify', sb=3, sa=3, fi=None)
        run = cp.add_run(txt); _sf(run, 13)
    return tbl

def add_footer_pagenum():
    """Thêm số trang vào footer section hiện tại"""
    sec = doc.sections[-1]
    sec.footer.is_linked_to_previous = False
    ftr = sec.footer
    for p in list(ftr.paragraphs):
        p._element.getparent().remove(p._element)
    fp = ftr.add_paragraph()
    _fmtp(fp, align='center', sb=0, sa=0, fi=None)
    run = fp.add_run(); _sf(run, 13)
    for tag, attr, val, txt in [
        ('w:fldChar',   'w:fldCharType', 'begin', None),
        ('w:instrText', None,            None,    ' PAGE '),
        ('w:fldChar',   'w:fldCharType', 'end',   None),
    ]:
        el = OxmlElement(tag)
        if attr: el.set(qn(attr), val)
        if txt:  el.set(qn('xml:space'), 'preserve'); el.text = txt
        run._r.append(el)

def add_section_break():
    """Thêm section break để bắt đầu vùng mới (có số trang)"""
    new_sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
    _cfg_sec(new_sec)
    return new_sec

# ─────────────────────────────────────────────
# PARSER MARKDOWN → DOCX
# ─────────────────────────────────────────────
def process_md(lines):
    i = 0
    while i < len(lines):
        line  = lines[i]
        strip = line.strip()

        # Bỏ divider, meta notes
        if strip in ('---',) or strip.startswith('> **Hệ thống:**') or \
           strip.startswith('> **Ngày') or strip.startswith('> **Công cụ') or \
           strip.startswith('> **Server'):
            i += 1; continue

        # Heading 1 (chương)
        if strip.startswith('# '):
            H1(strip[2:].strip()); i += 1; continue

        # Heading 2
        if strip.startswith('## '):
            H2(strip[3:].strip()); i += 1; continue

        # Heading 3
        if strip.startswith('### '):
            H3(strip[4:].strip()); i += 1; continue

        # Code block
        if strip.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i].rstrip()); i += 1
            if code_lines: CODE(code_lines)
            i += 1; continue

        # Bảng markdown
        if strip.startswith('|'):
            tbl_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                tbl_lines.append(lines[i].strip()); i += 1
            if len(tbl_lines) >= 2:
                def pr(r):
                    cols = [c.strip() for c in r.split('|')]
                    if cols and not cols[0]: cols = cols[1:]
                    if cols and not cols[-1]: cols = cols[:-1]
                    return cols
                hdrs = pr(tbl_lines[0])
                data = [pr(r) for r in tbl_lines[2:] if r.replace('|','').replace('-','').strip()]
                TABLE(hdrs, data)
            continue

        # Bullet -
        if strip.startswith('- '):
            LP(strip[2:].strip(), bullet=True); i += 1; continue

        # Numbered 1. 2. ...
        m = re.match(r'^(\d+)\.\s+(.+)$', strip)
        if m:
            NUMBERED(int(m.group(1)), m.group(2)); i += 1; continue

        # Blockquote >
        if strip.startswith('> '):
            QUOTE(strip[2:].strip()); i += 1; continue

        # Dòng trống
        if not strip:
            i += 1; continue

        # Đoạn thường
        P(strip); i += 1

# ─────────────────────────────────────────────
# ===  SECTION 0: TRANG BÌA CHÍNH  ===
# ─────────────────────────────────────────────
def cover_page_1():
    for txt, sz, bold in [
        ('ĐẠI HỌC THÁI NGUYÊN',                                       16, False),
        ('TRƯỜNG ĐẠI HỌC CÔNG NGHỆ THÔNG TIN VÀ TRUYỀN THÔNG',       14, True),
    ]:
        p = doc.add_paragraph()
        _fmtp(p, align='center', sb=6, sa=6, fi=None)
        r = p.add_run(txt); _sf(r, sz, bold)

    # Dòng trắng
    for _ in range(3): P()

    # Tên sinh viên
    p = doc.add_paragraph()
    _fmtp(p, align='center', sb=6, sa=6, fi=None)
    r = p.add_run('NGUYỄN XUÂN BẮC'); _sf(r, 16, True)

    P()

    # Tên đề tài
    p = doc.add_paragraph()
    _fmtp(p, align='center', sb=6, sa=6, fi=None)
    r = p.add_run('KIỂM THỬ HỆ THỐNG QUẢN LÝ KHÓA HỌC NGOẠI NGỮ TÍCH HỢP AI')
    _sf(r, 18, True)

    P()

    # Loại báo cáo
    p = doc.add_paragraph()
    _fmtp(p, align='center', sb=6, sa=6, fi=None)
    r = p.add_run('ĐỒ ÁN TỐT NGHIỆP ĐẠI HỌC'); _sf(r, 20, True)

    p = doc.add_paragraph()
    _fmtp(p, align='center', sb=6, sa=6, fi=None)
    r = p.add_run('CHUYÊN NGÀNH KỸ THUẬT PHẦN MỀM'); _sf(r, 14, True)

    for _ in range(5): P()

    p = doc.add_paragraph()
    _fmtp(p, align='center', sb=6, sa=6, fi=None)
    r = p.add_run('THÁI NGUYÊN, NĂM 2026'); _sf(r, 16, True)

cover_page_1()

# ─────────────────────────────────────────────
# ===  SECTION 1: BÌA LÓT  ===
# ─────────────────────────────────────────────
doc.add_page_break()

for txt, sz, bold in [
    ('ĐẠI HỌC THÁI NGUYÊN',                                       16, False),
    ('TRƯỜNG ĐẠI HỌC CÔNG NGHỆ THÔNG TIN VÀ TRUYỀN THÔNG',       14, True),
]:
    p = doc.add_paragraph()
    _fmtp(p, align='center', sb=6, sa=6, fi=None)
    r = p.add_run(txt); _sf(r, sz, bold)

for _ in range(2): P()

for txt, sz, bold in [
    ('ĐỒ ÁN',                    20, True),
    ('TỐT NGHIỆP ĐẠI HỌC',      20, True),
    ('CHUYÊN NGÀNH KỸ THUẬT PHẦN MỀM', 14, False),
]:
    p = doc.add_paragraph()
    _fmtp(p, align='center', sb=6, sa=6, fi=None)
    r = p.add_run(txt); _sf(r, sz, bold)

P()

p = doc.add_paragraph()
_fmtp(p, align='left', sb=6, sa=6, fi=None)
r = p.add_run('Đề tài:'); _sf(r, 16, False)

p = doc.add_paragraph()
_fmtp(p, align='center', sb=6, sa=6, fi=None)
r = p.add_run('KIỂM THỬ HỆ THỐNG QUẢN LÝ KHÓA HỌC NGOẠI NGỮ TÍCH HỢP AI')
_sf(r, 16, True)

P()

for label, value in [
    ('Sinh viên thực hiện: ', 'NGUYỄN XUÂN BẮC'),
    ('Lớp ',                  'KTPM K21'),
    ('Giáo viên hướng dẫn: ', 'ThS. NGUYỄN LAN OANH'),
]:
    p = doc.add_paragraph()
    _fmtp(p, align='center', sb=6, sa=6, fi=None)
    r1 = p.add_run(label);  _sf(r1, 16, False)
    r2 = p.add_run(value);  _sf(r2, 16, False)

for _ in range(4): P()

p = doc.add_paragraph()
_fmtp(p, align='center', sb=6, sa=6, fi=None)
r = p.add_run('THÁI NGUYÊN, NĂM 2026'); _sf(r, 16, True)

# ─────────────────────────────────────────────
# ===  LỜI CẢM ƠN  ===
# ─────────────────────────────────────────────
H1('LỜI CẢM ƠN')

for txt in [
    'Em xin chân thành bày tỏ lòng biết ơn sâu sắc đến quý thầy cô đã luôn '
    'tạo điều kiện thuận lợi và tận tình hỗ trợ em trong suốt quá trình học '
    'tập và thực hiện đề tài tốt nghiệp.',

    'Em đặc biệt xin gửi lời cảm ơn chân thành và sâu sắc nhất đến '
    'ThS. Nguyễn Lan Oanh – người đã trực tiếp hướng dẫn, truyền đạt kiến '
    'thức chuyên môn về kiểm thử phần mềm, định hướng nghiên cứu và góp ý '
    'quý báu giúp em hoàn thiện đề tài này.',

    'Em cũng xin chân thành cảm ơn Trường Đại học Công nghệ Thông tin và '
    'Truyền thông (ICTU) – Đại học Thái Nguyên, cùng toàn thể quý thầy cô '
    'trong Khoa Công nghệ phần mềm đã truyền đạt kiến thức và kỹ năng cần '
    'thiết trong suốt bốn năm học.',

    'Cuối cùng, em xin gửi lời cảm ơn đến gia đình và bạn bè đã luôn động '
    'viên, chia sẻ và tạo điều kiện tốt nhất để em hoàn thành đề tài.',

    'Mặc dù đã nỗ lực hết mình để hoàn thiện đề tài, em nhận thức rõ rằng '
    'công trình vẫn còn nhiều thiếu sót. Em kính mong nhận được sự chỉ bảo '
    'và góp ý của quý thầy cô để bài làm được hoàn chỉnh hơn.',

    'Một lần nữa, em xin chân thành cảm ơn quý thầy cô, cảm ơn cô Nguyễn '
    'Lan Oanh và toàn thể mọi người đã đồng hành cùng em.',
]:
    P(txt)

P()
SIGN_TABLE(
    '',
    'Thái Nguyên, tháng 05 năm 2026\nSinh viên thực hiện\n\n\n\nNguyễn Xuân Bắc'
)

# ─────────────────────────────────────────────
# ===  LỜI CAM ĐOAN  ===
# ─────────────────────────────────────────────
H1('LỜI CAM ĐOAN')

for txt in [
    'Em tên là: Nguyễn Xuân Bắc, sinh viên lớp: KTPM K21, trường: Đại học '
    'Công nghệ Thông tin và Truyền thông – Đại học Thái Nguyên.',

    'Em xin cam đoan rằng đề tài tốt nghiệp này là công trình nghiên cứu độc '
    'lập của cá nhân em dưới sự hướng dẫn của ThS. Nguyễn Lan Oanh. Toàn bộ '
    'nội dung bao gồm kế hoạch kiểm thử, thiết kế test case, thực thi kiểm '
    'thử tự động (Playwright), kiểm thử hiệu năng (Apache JMeter), kiểm thử '
    'bảo mật (OWASP ZAP) và phân tích mã nguồn (SonarQube) đều được em tự '
    'thực hiện.',

    'Các số liệu, kết quả kiểm thử được trình bày trong báo cáo là trung '
    'thực, khách quan và chưa từng được công bố trong bất kỳ công trình nào '
    'khác. Mọi thông tin tham khảo đều được trích dẫn và ghi rõ nguồn gốc.',

    'Em xin chịu hoàn toàn trách nhiệm trước quý thầy cô, Khoa Công nghệ '
    'Thông tin, và Nhà trường nếu phát hiện bất kỳ sự gian dối nào.',
]:
    P(txt)

P()
SIGN_TABLE(
    '',
    'Thái Nguyên, tháng 05 năm 2026\nNgười cam đoan\n\n\n\nNguyễn Xuân Bắc'
)

# ─────────────────────────────────────────────
# ===  MỤC LỤC  ===
# ─────────────────────────────────────────────
doc.add_page_break()
p = doc.add_paragraph(style='TOC Heading')
p.clear()
_fmtp(p, align='center', sb=6, sa=6, fi=None)
r = p.add_run('MỤC LỤC'); _sf(r, 13, True)

toc = [
    ('LỜI CẢM ƠN', 0),
    ('LỜI CAM ĐOAN', 0),
    ('MỤC LỤC', 0),
    ('MỤC LỤC HÌNH ẢNH', 0),
    ('MỤC LỤC BẢNG BIỂU', 0),
    ('LỜI MỞ ĐẦU', 0),
    ('CHƯƠNG 1: CƠ SỞ LÝ THUYẾT VỀ KIỂM THỬ PHẦN MỀM', 0),
    ('1.1 Tổng quan về kiểm thử phần mềm', 1),
    ('1.2 Các cấp độ kiểm thử', 1),
    ('1.3 Các loại kiểm thử theo kỹ thuật', 1),
    ('1.4 Kiểm thử tự động (Automated Testing)', 1),
    ('1.5 Kiểm thử bảo mật (Security Testing)', 1),
    ('1.6 CI/CD và tích hợp kiểm thử', 1),
    ('1.7 Các công cụ quản lý kiểm thử', 1),
    ('1.8 Mô hình kiểm thử trong dự án', 1),
    ('CHƯƠNG 2: KẾ HOẠCH VÀ THIẾT KẾ KIỂM THỬ HỆ THỐNG QUANLYKHOAHOC5', 0),
    ('2.1 Giới thiệu hệ thống', 1),
    ('2.2 Phạm vi kiểm thử', 1),
    ('2.3 Kế hoạch kiểm thử', 1),
    ('2.4 Thiết kế test case chi tiết', 1),
    ('2.5 Test data và điều kiện đầu', 1),
    ('2.6 Tiêu chí chấp nhận và metrics', 1),
    ('2.7 Cấu trúc thư mục KiemThu', 1),
    ('2.8 Rủi ro và biện pháp giảm thiểu', 1),
    ('CHƯƠNG 3: KẾT QUẢ KIỂM THỬ', 0),
    ('3.1 Tổng quan quá trình kiểm thử tự động', 1),
    ('3.2 Bảng tổng hợp kết quả 3 lần chạy', 1),
    ('3.3 Chi tiết kết quả từng module', 1),
    ('3.4 Tổng hợp lỗi và quá trình sửa chữa', 1),
    ('3.5 Kết quả kiểm thử hiệu năng – Apache JMeter', 1),
    ('3.6 Kết quả kiểm thử bảo mật – OWASP ZAP', 1),
    ('3.7 Kết quả phân tích mã nguồn – SonarQube', 1),
    ('3.8 Tích hợp CI/CD – GitHub Actions', 1),
    ('3.9 Kết luận tổng hợp', 1),
    ('KẾT LUẬN', 0),
    ('TÀI LIỆU THAM KHẢO', 0),
]
for entry, lvl in toc:
    p = doc.add_paragraph()
    _fmtp(p, align='left', sb=3, sa=3, fi=None, li=lvl*0.9)
    r = p.add_run(entry); _sf(r, 13, lvl == 0)

# ─────────────────────────────────────────────
# ===  MỤC LỤC HÌNH ẢNH  ===
# ─────────────────────────────────────────────
H1_nobreak('MỤC LỤC HÌNH ẢNH')
P('(Báo cáo sử dụng dữ liệu văn bản và bảng số liệu. Các biểu đồ kết quả '
  'kiểm thử được trình bày dưới dạng bảng trong nội dung tương ứng.)',
  italic=True)

# ─────────────────────────────────────────────
# ===  MỤC LỤC BẢNG BIỂU  ===
# ─────────────────────────────────────────────
H1_nobreak('MỤC LỤC BẢNG BIỂU')

TABLE(
    ['Số bảng', 'Tên bảng', 'Chương'],
    [
        ['Bảng 1.1', 'So sánh Playwright và Selenium WebDriver', 'Chương 1'],
        ['Bảng 1.2', 'Các chỉ số hiệu năng JMeter cần theo dõi', 'Chương 1'],
        ['Bảng 2.1', 'Thành phần công nghệ hệ thống QuanLyKhoaHoc5', 'Chương 2'],
        ['Bảng 2.2', 'Phân quyền người dùng – seed accounts', 'Chương 2'],
        ['Bảng 2.3', 'Phạm vi kiểm thử – In Scope / Out of Scope', 'Chương 2'],
        ['Bảng 2.4', 'Môi trường kiểm thử', 'Chương 2'],
        ['Bảng 2.5', 'Lịch kiểm thử 4 tuần', 'Chương 2'],
        ['Bảng 2.6', 'Bảng giá trị biên cho XepLoai điểm số', 'Chương 2'],
        ['Bảng 2.7', 'Yêu cầu phi chức năng và ngưỡng chấp nhận', 'Chương 2'],
        ['Bảng 2.8', 'Phân loại lỗi (Defect Severity)', 'Chương 2'],
        ['Bảng 2.9', 'Rủi ro và biện pháp giảm thiểu', 'Chương 2'],
        ['Bảng 3.1', 'Mục tiêu từng vòng chạy Playwright', 'Chương 3'],
        ['Bảng 3.2', 'Kết quả theo file test – 3 lần chạy', 'Chương 3'],
        ['Bảng 3.3', 'Thống kê tổng thể 3 lần chạy', 'Chương 3'],
        ['Bảng 3.4', 'Module 1 – Đăng nhập / Đăng xuất', 'Chương 3'],
        ['Bảng 3.5', 'Module 2 – Khóa học', 'Chương 3'],
        ['Bảng 3.6', 'Module 3 – Lớp học', 'Chương 3'],
        ['Bảng 3.7', 'Module 4 – Đăng ký khóa học', 'Chương 3'],
        ['Bảng 3.8', 'Module 5 – Điểm số', 'Chương 3'],
        ['Bảng 3.9', 'Module 6 – Thanh toán', 'Chương 3'],
        ['Bảng 3.10','Module 7 – Quản lý tài khoản', 'Chương 3'],
        ['Bảng 3.11','Module 8 – Phân công giảng viên', 'Chương 3'],
        ['Bảng 3.12','Tổng hợp lỗi script Lần 1', 'Chương 3'],
        ['Bảng 3.13','Tổng hợp lỗi script Lần 2', 'Chương 3'],
        ['Bảng 3.14','Lỗi còn tồn đọng sau Lần 3', 'Chương 3'],
        ['Bảng 3.15','Cấu hình kịch bản tải JMeter', 'Chương 3'],
        ['Bảng 3.16','Kết quả tổng hợp toàn bộ test JMeter', 'Chương 3'],
        ['Bảng 3.17','TG1 – Browse Public', 'Chương 3'],
        ['Bảng 3.18','TG2 – HocVien Auth Flow', 'Chương 3'],
        ['Bảng 3.19','TG3 – Admin Heavy Ops', 'Chương 3'],
        ['Bảng 3.20','Tóm tắt kết quả OWASP ZAP', 'Chương 3'],
        ['Bảng 3.21','Hai lần chạy SonarQube – so sánh', 'Chương 3'],
        ['Bảng 3.22','Kế hoạch cải thiện SonarQube', 'Chương 3'],
        ['Bảng 3.23','Tổng kết 4 công cụ kiểm thử', 'Chương 3'],
    ]
)

# ─────────────────────────────────────────────
# Thêm section break + footer số trang
# ─────────────────────────────────────────────
add_section_break()
add_footer_pagenum()

# ─────────────────────────────────────────────
# ===  LỜI MỞ ĐẦU  ===
# ─────────────────────────────────────────────
H1_nobreak('LỜI MỞ ĐẦU')

P('Lí do chọn đề tài', bold=True)
for txt in [
    'Trong quá trình phát triển phần mềm, kiểm thử đóng vai trò cực kỳ quan '
    'trọng trong việc đảm bảo chất lượng sản phẩm. Theo thống kê của NIST, '
    'chi phí sửa lỗi sau khi phát hành cao hơn từ 30 đến 100 lần so với khi '
    'phát hiện trong giai đoạn phát triển.',

    'Trong bối cảnh đó, kiểm thử tự động ngày càng được áp dụng rộng rãi '
    'nhằm nâng cao hiệu quả và độ phủ của quá trình kiểm thử. Hệ thống '
    'QuanLyKhoaHoc5 là ứng dụng web quản lý trung tâm ngoại ngữ tích hợp AI, '
    'được xây dựng trên nền tảng ASP.NET Core MVC (.NET 10). Hệ thống phục '
    'vụ ba nhóm người dùng với phân quyền rõ ràng và xử lý các nghiệp vụ '
    'quan trọng như đăng ký khóa học, quản lý điểm số, thanh toán học phí, '
    'phân công giảng dạy và gợi ý khóa học bằng AI.',

    'Chính vì những lý do trên, em đã chọn đề tài: "Kiểm thử hệ thống Quản '
    'lý khóa học ngoại ngữ tích hợp AI" nhằm áp dụng các kiến thức và kỹ '
    'năng kiểm thử phần mềm vào một hệ thống thực tế.',
]:
    P(txt)

P('Mục tiêu nghiên cứu', bold=True)
P('Đề tài hướng đến các mục tiêu chính sau:')
for txt in [
    'Hệ thống hóa cơ sở lý thuyết về kiểm thử phần mềm: các cấp độ, kỹ thuật, '
    'công cụ và quy trình theo chuẩn ISTQB/IEEE.',
    'Xây dựng kế hoạch và thiết kế 55 test case bao phủ 8 module nghiệp vụ '
    'chính của hệ thống QuanLyKhoaHoc5.',
    'Thực thi kiểm thử tự động với Playwright (E2E), kiểm thử hiệu năng với '
    'Apache JMeter (7.000 requests), kiểm thử bảo mật với OWASP ZAP và phân '
    'tích mã nguồn với SonarQube.',
    'Phân tích kết quả qua 3 vòng thực thi, ghi nhận lỗi, sửa script và đánh '
    'giá chất lượng toàn diện hệ thống.',
    'Đề xuất hướng khắc phục và hướng phát triển tiếp theo.',
]:
    LP(txt)

P('Bố cục báo cáo', bold=True)
P('Ngoài lời mở đầu, kết luận và tài liệu tham khảo, báo cáo gồm 3 chương:')
for txt in [
    '**Chương 1:** Cơ sở lý thuyết – trình bày các khái niệm nền tảng, cấp '
    'độ kiểm thử, kỹ thuật kiểm thử và các công cụ sử dụng.',
    '**Chương 2:** Kế hoạch và thiết kế kiểm thử – mô tả hệ thống, xác định '
    'phạm vi, thiết kế 55 test case chi tiết và xây dựng tiêu chí chấp nhận.',
    '**Chương 3:** Kết quả kiểm thử – trình bày kết quả thực thi 3 vòng '
    'Playwright, kết quả JMeter, ZAP và SonarQube với phân tích chi tiết.',
]:
    LP(txt)

# ─────────────────────────────────────────────
# ===  CHƯƠNG 1  ===
# ─────────────────────────────────────────────
with open(os.path.join(BASE, 'Chuong1_Co_So_Ly_Thuyet.md'), encoding='utf-8') as f:
    process_md(f.readlines())

# ─────────────────────────────────────────────
# ===  CHƯƠNG 2  ===
# ─────────────────────────────────────────────
with open(os.path.join(BASE, 'Chuong2_Ke_Hoach_Va_Thiet_Ke.md'), encoding='utf-8') as f:
    process_md(f.readlines())

# ─────────────────────────────────────────────
# ===  CHƯƠNG 3  ===
# ─────────────────────────────────────────────
with open(os.path.join(BASE, 'Chuong3_Ket_Qua_Kiem_Thu.md'), encoding='utf-8') as f:
    process_md(f.readlines())

# ─────────────────────────────────────────────
# ===  KẾT LUẬN  ===
# ─────────────────────────────────────────────
H1('KẾT LUẬN')

P('* Kết quả đạt được', bold=True)
for txt in [
    'Về kiến thức: Em đã nắm vững cơ sở lý thuyết về kiểm thử phần mềm theo '
    'chuẩn ISTQB, hiểu rõ các cấp độ kiểm thử, kỹ thuật hộp đen/trắng/xám, '
    'và vận dụng thành thạo 4 công cụ kiểm thử hiện đại.',
    'Về kết quả Playwright (E2E): Hoàn thành 3 vòng kiểm thử với 55 test case '
    'trên 8 module. Lần 3 đạt tỷ lệ pass 96.7% (58/60 TCs); 6/8 module đạt '
    '100%. Thời gian thực thi giảm 47% nhờ phương pháp kiểm thử lặp.',
    'Về hiệu năng JMeter: 7.000 requests, 450 VUsers, response time trung bình '
    '2ms, tối đa 193ms, throughput 53.6 req/s. Error rate 17.86% được xác nhận '
    'là do thiếu CSRF token trong cấu hình JMeter.',
    'Về bảo mật ZAP: 0 High Risk. 3 cảnh báo Medium đều là HTTP headers thiếu, '
    'có thể khắc phục bằng 1 middleware. Các cơ chế cốt lõi (EF Core, '
    'AntiForgery, BCrypt, AuthorizeRole) hoạt động đúng.',
    'Về SonarQube: Phát hiện 477 issues/20k LOC; Coverage 0.0%; Duplication '
    '9.57% → Quality Gate Failed. Đây là technical debt cần đầu tư dài hạn.',
]:
    LP(txt)

P('* Hạn chế', bold=True)
for txt in [
    'Script Playwright chưa bao phủ đầy đủ: Module 4 chỉ cover 3/7 TC; '
    'Module 5 chỉ thực thi 1/6 TC thực chất.',
    'JMeter thiếu CSRF token extraction → không thể đánh giá authenticated flow.',
    'Coverage SonarQube = 0.0% vì không có unit test ở tầng Service/Repository.',
    'CI/CD chưa kích hoạt do repository chưa push lên GitHub remote.',
]:
    LP(txt)

P('* Hướng phát triển', bold=True)
for txt in [
    'Ngắn hạn: Sửa dialog handler trong `04_dang_ky_test.js`, thêm CSRF '
    'extraction vào JMX, thêm security headers middleware vào Program.cs.',
    'Trung hạn: Viết xUnit unit tests đạt Coverage ≥ 20%; tạo RoleConstants và '
    'BaseController để giảm duplication từ 9.57% xuống < 3%.',
    'Dài hạn: Đạt Quality Gate Green; tích hợp CI/CD hoàn chỉnh; mở rộng '
    'kiểm thử sang mobile browser và Mutation Testing.',
]:
    LP(txt)

# ─────────────────────────────────────────────
# ===  TÀI LIỆU THAM KHẢO  ===
# ─────────────────────────────────────────────
H1('TÀI LIỆU THAM KHẢO')

for ref in [
    '[1] ISTQB Việt Nam, Tài liệu chứng chỉ kiểm thử phần mềm ISTQB Foundation Level v4.0. '
    'https://www.istqb.org/certifications/certified-tester-foundation-level',
    '[2] Nguyễn Lan Oanh (2024). Bài giảng Kiểm thử phần mềm. Trường ĐH Công nghệ '
    'Thông tin và Truyền thông – ĐH Thái Nguyên.',
    '[3] Myers, G. J., Sandler, C., & Badgett, T. (2012). The Art of Software Testing '
    '(3rd ed.). Wiley.',
    '[4] IEEE Std 829-2008. IEEE Standard for Software and System Test Documentation. '
    'IEEE Computer Society.',
    '[5] Microsoft (2024). Playwright for .NET Documentation. '
    'https://playwright.dev/dotnet/docs/intro',
    '[6] Apache Software Foundation (2023). Apache JMeter User Manual v5.6. '
    'https://jmeter.apache.org/usermanual/index.html',
    '[7] OWASP Foundation (2021). OWASP Top Ten 2021. '
    'https://owasp.org/www-project-top-ten/',
    '[8] SonarSource (2024). SonarQube Documentation – Community Edition. '
    'https://docs.sonarsource.com/sonarqube/latest/',
    '[9] Microsoft (2024). ASP.NET Core MVC in .NET 10 – Security and Authentication. '
    'https://learn.microsoft.com/en-us/aspnet/core/security/',
    '[10] NIST (2002). The Economic Impacts of Inadequate Infrastructure for Software '
    'Testing. National Institute of Standards and Technology.',
]:
    P(ref)

# ─────────────────────────────────────────────
# LƯU FILE
# ─────────────────────────────────────────────
# Dọn các script và file phân tích tạm
for tmp in [
    r'D:\QuanLyKhoaHoc5\analyze_template.py',
    r'D:\QuanLyKhoaHoc5\check_footer.py',
    r'D:\QuanLyKhoaHoc5\template_analysis.txt',
]:
    try: os.remove(tmp)
    except: pass

doc.save(OUT)
print(f'OK  {OUT}')
print(f'    {os.path.getsize(OUT):,} bytes')
